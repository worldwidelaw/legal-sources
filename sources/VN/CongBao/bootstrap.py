#!/usr/bin/env python3
"""
Vietnam Official Gazette (Công Báo) Data Fetcher

Fetches Vietnamese legislation from congbao.chinhphu.vn, the official
electronic gazette published by the Office of the Government.

Approach:
  - Listing: paginate /van-ban-dang-cong-bao/trang-{page}.htm (10 per page)
  - Detail: fetch each document page for metadata + DOCX download URL
  - Full text: download DOCX files and extract text via python-docx

Covers ~25,000 documents: laws, decrees, decisions, circulars, ordinances,
resolutions, consolidated texts.
"""

import html as html_mod
import json
import logging
import os
import re
import sys
import tempfile
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

BASE_URL = "https://congbao.chinhphu.vn"
LISTING_URL = BASE_URL + "/van-ban-dang-cong-bao/trang-{page}.htm"
CDN_DOWNLOAD = "https://g7.cdnchinhphu.vn/api/download/stream"
DELAY = 1.5
HEADERS = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) LegalDataHunter/1.0"}

try:
    import urllib.request
    import urllib.parse
except ImportError:
    pass


def http_get(url: str, timeout: int = 30) -> Optional[str]:
    """Fetch a URL and return decoded text, or None on failure."""
    req = urllib.request.Request(url, headers=HEADERS)
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return resp.read().decode("utf-8", errors="replace")
    except Exception as e:
        logger.warning(f"HTTP GET failed for {url[:120]}: {e}")
        return None


def http_download(url: str, timeout: int = 60) -> Optional[bytes]:
    """Download binary content from a URL."""
    req = urllib.request.Request(url, headers=HEADERS)
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return resp.read()
    except Exception as e:
        logger.warning(f"Download failed for {url[:120]}: {e}")
        return None


def parse_vn_date(date_str: str) -> Optional[str]:
    """Parse Vietnamese date DD/MM/YYYY to ISO 8601."""
    if not date_str:
        return None
    date_str = date_str.strip()
    for fmt in ("%d/%m/%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(date_str, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return None


def extract_doc_number(title: str) -> Optional[str]:
    """Extract document number from title, e.g. '17/2026/QĐ-TTg' from 'Quyết định số 17/2026/QĐ-TTg ...'."""
    m = re.search(r"số\s+([\w/.-]+)", title, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    return None


def classify_doc_type(title: str, slug: str) -> str:
    """Classify document type from title or URL slug."""
    title_lower = title.lower()
    if "luật" in title_lower or slug.startswith("luat-"):
        return "Luật"
    elif "nghị định" in title_lower or "nghi-dinh" in slug:
        return "Nghị định"
    elif "nghị quyết" in title_lower or "nghi-quyet" in slug:
        return "Nghị quyết"
    elif "quyết định" in title_lower or "quyet-dinh" in slug:
        return "Quyết định"
    elif "thông tư" in title_lower or "thong-tu" in slug:
        return "Thông tư"
    elif "pháp lệnh" in title_lower or "phap-lenh" in slug:
        return "Pháp lệnh"
    elif "chỉ thị" in title_lower or "chi-thi" in slug:
        return "Chỉ thị"
    elif "văn bản hợp nhất" in title_lower or "van-ban-hop-nhat" in slug:
        return "Văn bản hợp nhất"
    return "Văn bản"


def extract_text_from_docx(content: bytes) -> Optional[str]:
    """Extract text from DOCX binary content using python-docx, with ZIP fallback."""
    try:
        import docx
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return None

    import io

    # Try python-docx first
    try:
        doc = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        text = "\n".join(paragraphs)
        if text.strip() and len(text) > 50:
            return text.strip()
    except Exception as e:
        logger.debug(f"python-docx failed: {e}, trying XML extraction")

    # Fallback: extract text directly from document.xml inside the ZIP
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            # Find the main document XML (handle both / and \ separators)
            doc_xml = None
            for name in z.namelist():
                normalized = name.replace("\\", "/")
                if normalized.endswith("word/document.xml") or normalized == "word/document.xml":
                    doc_xml = name
                    break
            if not doc_xml:
                # Try any XML file that might contain document content
                for name in z.namelist():
                    normalized = name.replace("\\", "/")
                    if "document" in normalized.lower() and normalized.endswith(".xml"):
                        doc_xml = name
                        break
            if doc_xml:
                xml_content = z.read(doc_xml).decode("utf-8", errors="replace")
                # Strip XML tags to get plain text
                text = re.sub(r"<[^>]+>", " ", xml_content)
                text = html_mod.unescape(text)
                text = re.sub(r"\s+", " ", text).strip()
                # Split on common separators
                text = re.sub(r"\s{3,}", "\n", text)
                if len(text) > 50:
                    return text
    except Exception as e:
        logger.warning(f"ZIP XML extraction failed: {e}")

    return None


class CongBaoFetcher:
    """Fetcher for Vietnamese legislation from congbao.chinhphu.vn."""

    def __init__(self):
        self.delay = DELAY

    def get_total_pages(self) -> int:
        """Get total number of listing pages.

        Returns -1 when the listing root could not be fetched at all (network
        block / connection reset — typical of a datacenter-IP block on VN gov
        hosts) so callers can distinguish "host unreachable" from a genuine
        "0 pages". Returns 0 only when the page loaded but the pagination token
        is absent (real upstream markup change worth investigating).
        """
        url = BASE_URL + "/van-ban-dang-cong-bao.htm"
        data = http_get(url)
        if not data:
            logger.error(
                "Listing root %s unreachable (empty response) — likely a "
                "datacenter-IP block; cannot determine page count", url
            )
            return -1
        m = re.search(r"totalPageSodo\s*=\s*(\d+)", data)
        if not m:
            logger.error(
                "Pagination token 'totalPageSodo' not found on %s — upstream "
                "markup may have changed", url
            )
            return 0
        return int(m.group(1))

    def list_documents(self, page: int) -> List[Dict[str, str]]:
        """Fetch one listing page and return document URLs with titles."""
        url = LISTING_URL.format(page=page)
        data = http_get(url)
        if not data:
            return []

        results = []
        seen = set()
        for m in re.finditer(
            r'<a[^>]*href="(/van-ban/([^"]+)-(\d+)\.htm)"[^>]*title="([^"]*)"',
            data,
        ):
            path, slug, doc_id, title = m.group(1), m.group(2), m.group(3), m.group(4)
            if doc_id in seen:
                continue
            seen.add(doc_id)
            results.append({
                "path": path,
                "slug": slug,
                "doc_id": doc_id,
                "title": html_mod.unescape(title),
            })
        return results

    def fetch_document_page(self, path: str) -> Optional[Dict[str, Any]]:
        """Fetch a document detail page and extract metadata + DOCX URL."""
        url = BASE_URL + path
        data = http_get(url, timeout=30)
        if not data:
            return None

        # Title from OG meta or h1
        title = None
        og_m = re.search(r'property="og:title"[^>]*content="([^"]+)"', data)
        if og_m:
            title = html_mod.unescape(og_m.group(1)).strip()
        if not title:
            h1_m = re.search(r'class="title"[^>]*>(.*?)</h1>', data, re.DOTALL)
            if h1_m:
                title = html_mod.unescape(re.sub(r"<[^>]+>", "", h1_m.group(1))).strip()

        # Issue date: "Ban hành: DD/MM/YYYY"
        issue_date = None
        date_m = re.search(r"Ban hành:\s*(\d{2}/\d{2}/\d{4})", data)
        if date_m:
            issue_date = parse_vn_date(date_m.group(1))

        # Effective date: "Hiệu lực: DD/MM/YYYY"
        effective_date = None
        eff_m = re.search(r"Hiệu lực:\s*(\d{2}/\d{2}/\d{4})", data)
        if eff_m:
            effective_date = parse_vn_date(eff_m.group(1))

        # DOCX download URL (preferred over PDF for text extraction)
        docx_url = None
        docx_m = re.search(
            r'href="(https://g7\.cdnchinhphu\.vn/api/download/stream[^"]*\.docx[^"]*)"',
            data,
        )
        if docx_m:
            docx_url = html_mod.unescape(docx_m.group(1))

        # Always extract PDF URL as fallback
        pdf_url = None
        pdf_m = re.search(
            r'href="(https://g7\.cdnchinhphu\.vn/api/download/stream[^"]*\.pdf[^"]*)"',
            data,
        )
        if pdf_m:
            pdf_url = html_mod.unescape(pdf_m.group(1))

        # Gazette number
        gazette_num = None
        gz_m = re.search(r'Công báo số\s*(\d+)', data)
        if gz_m:
            gazette_num = gz_m.group(1)

        return {
            "title": title,
            "issue_date": issue_date,
            "effective_date": effective_date,
            "docx_url": docx_url,
            "pdf_url": pdf_url,
            "gazette_number": gazette_num,
            "page_url": url,
        }

    def fetch_full_text(self, docx_url: Optional[str], pdf_url: Optional[str]) -> Optional[str]:
        """Download and extract full text from DOCX (preferred) or PDF."""
        if docx_url:
            content = http_download(docx_url, timeout=60)
            if content:
                text = extract_text_from_docx(content)
                if text:
                    return text
                logger.warning("DOCX extraction failed, trying PDF fallback")

        if pdf_url:
            content = http_download(pdf_url, timeout=60)
            if content:
                try:
                    from common.pdf_extract import extract_pdf_markdown
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                        f.write(content)
                        tmp_path = f.name
                    try:
                        text = extract_pdf_markdown(tmp_path)
                        return text.strip() if text and len(text) > 50 else None
                    finally:
                        os.unlink(tmp_path)
                except Exception as e:
                    logger.warning(f"PDF extraction failed: {e}")

        return None

    def fetch_document(self, item: Dict[str, str]) -> Optional[Dict[str, Any]]:
        """Fetch a complete document with full text."""
        path = item["path"]
        doc_id = item["doc_id"]
        slug = item["slug"]
        list_title = item["title"]

        page_data = self.fetch_document_page(path)
        if not page_data:
            return None

        title = page_data["title"] or list_title
        if not title:
            return None

        time.sleep(self.delay)

        # Download full text
        text = self.fetch_full_text(page_data["docx_url"], page_data["pdf_url"])
        if not text:
            logger.warning(f"No full text for doc_id={doc_id}: {title[:60]}")
            return None

        doc_number = extract_doc_number(title)
        doc_type = classify_doc_type(title, slug)

        return {
            "_id": f"VN-CongBao-{doc_id}",
            "_source": "VN/CongBao",
            "_type": "legislation",
            "_fetched_at": datetime.now(timezone.utc).isoformat(),
            "title": title,
            "text": text,
            "date": page_data["issue_date"],
            "url": page_data["page_url"],
            "document_number": doc_number,
            "document_type": doc_type,
            "effective_date": page_data["effective_date"],
            "gazette_number": page_data["gazette_number"],
        }

    def normalize(self, raw: Dict[str, Any]) -> Dict[str, Any]:
        """Already normalized during fetch."""
        return raw

    def fetch_all(self) -> Iterator[Dict[str, Any]]:
        """Yield all documents from congbao.chinhphu.vn.

        Discovery does not depend solely on the ``totalPageSodo`` token: even
        when it is missing (markup change) we still walk the paginated
        ``trang-{page}.htm`` listing until a page comes back empty. The token,
        when present, only supplies an upper bound / progress denominator.
        """
        total_pages = self.get_total_pages()
        if total_pages < 0:
            # Listing root unreachable — surface it loudly rather than silently
            # yielding nothing (which the fleet reads as "0 records, complete").
            raise RuntimeError(
                "congbao.chinhphu.vn listing root is unreachable "
                "(datacenter-IP block?) — aborting so the run fails loud "
                "instead of reporting a false-empty corpus"
            )

        # Cap: known total when the token parsed, otherwise a generous ceiling.
        # The loop stops early on the first empty page regardless.
        page_cap = total_pages if total_pages > 0 else 5000
        logger.info(f"Total listing pages: {total_pages} (walking up to {page_cap})")

        empty_streak = 0
        for page in range(1, page_cap + 1):
            logger.info(f"Listing page {page}/{page_cap}")
            items = self.list_documents(page)
            if not items:
                # Tolerate a single transient empty page, stop on two in a row.
                empty_streak += 1
                logger.warning(f"Empty page {page} (streak {empty_streak})")
                if empty_streak >= 2:
                    logger.info("Two consecutive empty pages — end of listing")
                    break
                time.sleep(self.delay)
                continue
            empty_streak = 0

            for item in items:
                doc = self.fetch_document(item)
                if doc:
                    yield doc
                time.sleep(self.delay)

    def fetch_updates(self, since: str) -> Iterator[Dict[str, Any]]:
        """Fetch documents published since a given date."""
        since_dt = datetime.fromisoformat(since)
        for page in range(1, 200):
            items = self.list_documents(page)
            if not items:
                break
            all_old = True
            for item in items:
                doc = self.fetch_document(item)
                if doc and doc.get("date"):
                    try:
                        doc_dt = datetime.fromisoformat(doc["date"])
                        if doc_dt >= since_dt:
                            all_old = False
                            yield doc
                    except (ValueError, TypeError):
                        yield doc
                elif doc:
                    yield doc
                time.sleep(self.delay)
            if all_old:
                break


def bootstrap_sample(sample_dir: Path, count: int = 15):
    """Fetch sample documents and save to sample directory."""
    sample_dir.mkdir(parents=True, exist_ok=True)
    fetcher = CongBaoFetcher()

    total_pages = fetcher.get_total_pages()
    logger.info(f"congbao.chinhphu.vn has {total_pages} listing pages (~{total_pages * 10} documents)")

    items = fetcher.list_documents(1)
    logger.info(f"Got {len(items)} items from page 1")

    saved = 0
    for item in items[:count]:
        doc_id = item["doc_id"]
        logger.info(f"Fetching document {doc_id}: {item['title'][:60]}")

        doc = fetcher.fetch_document(item)
        if not doc:
            logger.warning(f"Skipping doc_id={doc_id} (no content)")
            continue

        text_len = len(doc.get("text", ""))
        logger.info(f"  Title: {doc.get('title', 'N/A')[:80]}")
        logger.info(f"  Text: {text_len} chars")
        logger.info(f"  Date: {doc.get('date')}")

        out_file = sample_dir / f"{doc['_id']}.json"
        with open(out_file, "w", encoding="utf-8") as f:
            json.dump(doc, f, ensure_ascii=False, indent=2)

        saved += 1
        logger.info(f"  Saved ({saved}/{count})")

    logger.info(f"Bootstrap complete: {saved} documents saved to {sample_dir}")
    return saved


if __name__ == "__main__":
    source_dir = Path(__file__).parent
    sample_dir = source_dir / "sample"

    if len(sys.argv) > 1 and sys.argv[1] == "bootstrap":
        sample_flag = "--sample" in sys.argv
        count = 15 if sample_flag else 50
        saved = bootstrap_sample(sample_dir, count)
        if saved < 10:
            logger.error(f"Only {saved} documents saved, expected at least 10")
            sys.exit(1)
    else:
        print("Usage: python3 bootstrap.py bootstrap [--sample]")
        print("  bootstrap --sample  Fetch 15 sample documents")
        print("  bootstrap           Fetch 50 sample documents")
