#!/usr/bin/env python3
"""
HR/ConstitutionalCourt -- Croatian Constitutional Court Case Law Fetcher

Fetches case law from the Croatian Constitutional Court (Ustavni sud)
decision database at sljeme.usud.hr.

Strategy:
  - IBM Domino/XPages platform with session-based navigation
  - Browse-by-date view to enumerate all decisions by year and case type
  - XPages partial refresh POST requests to expand year/category nodes
  - Individual decision pages provide metadata + DOCX/PDF download links
  - Full text extracted from DOCX attachments via python-docx

Endpoints:
  - Browse: https://sljeme.usud.hr/usud/praksaw.nsf/vPremaDatumuDonos.xsp
  - Decision: https://sljeme.usud.hr/usud/praksaw.nsf/fOdluka.xsp?action=openDocument&documentId={UNID}
  - DOCX: https://sljeme.usud.hr/Usud/Praksaw.nsf/{UNID}/$FILE/{filename}.docx

Data:
  - ~10,000-25,000 decisions (2007-present)
  - Case types: U-I through U-X, SuP-O, SuS, SuT-A
  - Language: Croatian (HRV)
  - Rate limit: ~2 seconds between requests

Usage:
  python bootstrap.py bootstrap          # Full initial pull
  python bootstrap.py bootstrap --sample # Fetch 10+ sample records
  python bootstrap.py update             # Incremental update
  python bootstrap.py test               # Quick connectivity test
"""

import sys
import json
import logging
import re
import html
import time
import tempfile
import os
from pathlib import Path
from datetime import datetime, timezone
from typing import Generator, Optional, Dict, Any, List, Tuple
from io import BytesIO

# Add project root to path
PROJECT_ROOT = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(PROJECT_ROOT))

from common.base_scraper import BaseScraper
from common.http_client import HttpClient

from common.pdf_extract import extract_pdf_markdown


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("legal-data-hunter.HR.ConstitutionalCourt")

BASE_URL = "https://sljeme.usud.hr"
BROWSE_PATH = "/usud/praksaw.nsf/vPremaDatumuDonos.xsp"
DECISION_PATH = "/usud/praksaw.nsf/fOdluka.xsp"
NSF_PATH = "/Usud/Praksaw.nsf"


class CroatianConstitutionalCourtScraper(BaseScraper):
    """
    Scraper for HR/ConstitutionalCourt -- Croatian Constitutional Court.
    Country: HR
    URL: https://sljeme.usud.hr

    Data types: case_law
    Auth: none (Open Government Data)
    """

    def __init__(self):
        source_dir = Path(__file__).parent
        super().__init__(source_dir)

        self.client = HttpClient(
            base_url=BASE_URL,
            headers={
                "User-Agent": "LegalDataHunter/1.0 (Open Data Research)",
                "Accept": "text/html,application/xhtml+xml",
                "Accept-Language": "hr,en",
            },
            timeout=60,
        )
        self._viewid = None

    def _get_session(self) -> str:
        """Initialize session and get $$viewid from the browse page."""
        resp = self.client.get(BROWSE_PATH)
        resp.raise_for_status()

        viewid_match = re.search(
            r'name="\$\$viewid"[^>]*value="([^"]+)"', resp.text
        )
        if not viewid_match:
            raise RuntimeError("Could not extract $$viewid from browse page")

        self._viewid = viewid_match.group(1)
        logger.info(f"Session initialized, viewid: {self._viewid}")
        return resp.text

    def _xsp_post(self, submit_id: str) -> str:
        """Send an XPages partial refresh POST to expand a tree node."""
        post_data = {
            "$$viewid": self._viewid,
            "$$xspsubmitid": submit_id,
            "$$xspexecid": submit_id,
            "$$xspsubmitvalue": "",
            "$$xspsubmitscroll": "0|0",
            "view:_id1": "view:_id1",
        }
        headers = {
            "Content-Type": "application/x-www-form-urlencoded",
            "X-Requested-With": "XMLHttpRequest",
            "Referer": f"{BASE_URL}{BROWSE_PATH}",
        }

        self.rate_limiter.wait()
        resp = self.client.session.post(
            f"{BASE_URL}{BROWSE_PATH}",
            data=post_data,
            headers=headers,
            timeout=60,
        )
        resp.raise_for_status()

        # Update viewid from response
        new_viewid = re.search(
            r'name="\$\$viewid"[^>]*value="([^"]+)"', resp.text
        )
        if new_viewid:
            self._viewid = new_viewid.group(1)

        return resp.text

    def _enumerate_years(self, page_html: str) -> List[Tuple[int, str]]:
        """Extract year indices and their expand trigger IDs from the browse page."""
        years = []
        # Year expand triggers follow pattern: rpView:{index}:_id45:_id46
        triggers = re.findall(
            r'XSP\.attachPartial\([^)]*?"(view:_id1:_id2:_id3:callbackContent:fctView:_id38:rpView:(\d+):_id45:_id46)"',
            page_html,
        )
        # Also extract year labels
        year_labels = re.findall(
            r'rpView:(\d+):_id45:lblNazivKategorije"[^>]*>(\d{4})<', page_html
        )
        year_map = {idx: int(year) for idx, year in year_labels}

        for trigger_id, idx in triggers:
            year = year_map.get(idx, 2026 - int(idx))
            years.append((year, trigger_id))

        return sorted(years, key=lambda x: -x[0])  # newest first

    def _enumerate_subcategories(
        self, year_html: str, year_idx: int
    ) -> List[Tuple[str, str]]:
        """Extract subcategory (case type) names and trigger IDs after expanding a year."""
        subcats = []
        # Category triggers: rpView:{year_idx}:_id45:facetContent:rpCategory:{cat_idx}:_id49:_id50
        pattern = (
            rf'rpView:{year_idx}:_id45:facetContent:rpCategory:(\d+):_id49:_id50'
        )
        cat_triggers = re.findall(
            rf'XSP\.attachPartial\([^)]*?"(view:_id1:_id2:_id3:callbackContent:fctView:_id38:{pattern})"',
            year_html,
        )

        # Extract category labels
        label_pattern = (
            rf'rpView:{year_idx}:_id45:facetContent:rpCategory:(\d+):_id49:lblNazivKategorije"[^>]*>([^<]+)<'
        )
        labels = re.findall(label_pattern, year_html)
        label_map = {idx: name for idx, name in labels}

        for trigger_id, cat_idx in cat_triggers:
            name = label_map.get(cat_idx, f"category-{cat_idx}")
            subcats.append((name, trigger_id))

        return subcats

    def _extract_document_links(
        self, html_content: str
    ) -> List[Dict[str, str]]:
        """Extract document IDs and inline metadata from expanded category HTML."""
        docs = []
        # Pattern: href="...documentId={UNID}">{case_number} - {date} - {type} - {outcome}</a>
        pattern = re.compile(
            r'documentId=([A-F0-9]{32})"[^>]*>\s*'
            r'([^<]+)<',
            re.IGNORECASE,
        )
        for match in pattern.finditer(html_content):
            unid = match.group(1)
            link_text = html.unescape(match.group(2).strip())

            # Parse inline metadata: "SuT-A-16/2024 - 25.06.2024 - odluka - ostalo"
            parts = [p.strip() for p in link_text.split(" - ")]
            doc = {"unid": unid, "link_text": link_text}
            if len(parts) >= 1:
                doc["signatura"] = parts[0]
            if len(parts) >= 2:
                doc["date_str"] = parts[1]
            if len(parts) >= 3:
                doc["decision_type"] = parts[2]
            if len(parts) >= 4:
                doc["outcome"] = parts[3]

            docs.append(doc)

        return docs

    def _fetch_decision_page(self, unid: str) -> Dict[str, Any]:
        """Fetch the decision page for metadata and file download URLs."""
        self.rate_limiter.wait()
        url = f"{DECISION_PATH}?action=openDocument&documentId={unid}"
        resp = self.client.get(url)
        resp.raise_for_status()

        content = resp.text
        meta = {"unid": unid}

        # Extract structured metadata fields
        field_patterns = {
            "signatura": r'cfSignatura"[^>]*>([^<]+)',
            "outcome": r'cfOdluka"[^>]*>([^<]+)',
            "decision_date": r'cfDatum"[^>]*>([^<]+)',
            "decision_type": r'cfVrstaOdluke"[^>]*>([^<]+)',
            "summary": r'cfZakljucak"[^>]*>([^<]+)',
        }
        for key, pattern in field_patterns.items():
            m = re.search(pattern, content)
            if m:
                meta[key] = html.unescape(m.group(1).strip())

        # Extract DOCX and PDF download URLs
        file_urls = re.findall(
            r'href="(https?://[^"]*\$FILE[^"]*\.(?:docx?|pdf))"',
            content,
            re.IGNORECASE,
        )
        if not file_urls:
            # Try without full URL
            file_urls = re.findall(
                r'href="([^"]*\$FILE[^"]*\.(?:docx?|pdf))"',
                content,
                re.IGNORECASE,
            )
            # Also try direct file links
            if not file_urls:
                file_urls = re.findall(
                    r'href="([^"]*Praksaw\.nsf[^"]*\.(?:docx?|pdf))"',
                    content,
                    re.IGNORECASE,
                )

        meta["docx_urls"] = [
            u for u in file_urls if u.lower().endswith(".docx")
        ]
        meta["pdf_urls"] = [
            u for u in file_urls if u.lower().endswith(".pdf")
        ]

        return meta

    def _download_full_text(self, meta: Dict[str, Any]) -> str:
        """Download DOCX (preferred) or PDF and extract full text."""
        import docx as python_docx

        # Try DOCX first
        for url in meta.get("docx_urls", []):
            try:
                self.rate_limiter.wait()
                if not url.startswith("http"):
                    url = f"{BASE_URL}{url}"
                resp = self.client.session.get(url, timeout=60)
                if resp.status_code == 200 and len(resp.content) > 100:
                    doc = python_docx.Document(BytesIO(resp.content))
                    text = "\n".join(
                        p.text for p in doc.paragraphs if p.text.strip()
                    )
                    if text.strip():
                        return text.strip()
            except Exception as e:
                logger.debug(f"DOCX download failed for {url}: {e}")

        # Fallback to PDF
        if pypdf:
            for url in meta.get("pdf_urls", []):
                try:
                    self.rate_limiter.wait()
                    if not url.startswith("http"):
                        url = f"{BASE_URL}{url}"
                    resp = self.client.session.get(url, timeout=60)
                    if resp.status_code == 200 and len(resp.content) > 100:
                        reader = pypdf.PdfReader(BytesIO(resp.content))
                        text = "\n".join(
                            page.extract_text() or ""
                            for page in reader.pages
                        )
                        if text.strip():
                            return text.strip()
                except Exception as e:
                    logger.debug(f"PDF download failed for {url}: {e}")

        # Last resort: construct URL from signatura
        signatura = meta.get("signatura", "")
        unid = meta.get("unid", "")
        if signatura and unid:
            filename = signatura.replace("/", "-")
            for ext in [".docx", ".pdf"]:
                try:
                    url = f"{BASE_URL}{NSF_PATH}/{unid}/%24FILE/{filename}{ext}"
                    self.rate_limiter.wait()
                    resp = self.client.session.get(url, timeout=60)
                    if resp.status_code == 200 and len(resp.content) > 100:
                        if ext == ".docx":
                            doc = python_docx.Document(BytesIO(resp.content))
                            text = "\n".join(
                                p.text
                                for p in doc.paragraphs
                                if p.text.strip()
                            )
                        else:
                            if pypdf:
                                reader = pypdf.PdfReader(
                                    BytesIO(resp.content)
                                )
                                text = "\n".join(
                                    page.extract_text() or ""
                                    for page in reader.pages
                                )
                            else:
                                continue
                        if text.strip():
                            return text.strip()
                except Exception as e:
                    logger.debug(f"Constructed URL failed: {e}")

        return ""

    def _enumerate_all_documents(
        self, sample_mode: bool = False, sample_size: int = 12
    ) -> Generator[Dict[str, Any], None, None]:
        """
        Enumerate all documents by expanding year/category tree nodes.
        Yields raw document dicts with UNID and inline metadata.
        """
        page_html = self._get_session()
        years = self._enumerate_years(page_html)
        logger.info(f"Found {len(years)} years: {[y[0] for y in years]}")

        total_docs = 0

        for year, year_trigger in years:
            if sample_mode and total_docs >= sample_size:
                break

            logger.info(f"Expanding year {year}...")
            try:
                year_html = self._xsp_post(year_trigger)
            except Exception as e:
                logger.warning(f"Failed to expand year {year}: {e}")
                continue

            # Find the year index from the trigger
            idx_match = re.search(r'rpView:(\d+):', year_trigger)
            year_idx = idx_match.group(1) if idx_match else "0"

            subcats = self._enumerate_subcategories(year_html, year_idx)
            logger.info(
                f"  Year {year}: {len(subcats)} subcategories: {[s[0] for s in subcats]}"
            )

            for cat_name, cat_trigger in subcats:
                if sample_mode and total_docs >= sample_size:
                    break

                try:
                    cat_html = self._xsp_post(cat_trigger)
                except Exception as e:
                    logger.warning(
                        f"Failed to expand {year}/{cat_name}: {e}"
                    )
                    continue

                docs = self._extract_document_links(cat_html)
                logger.info(
                    f"  {year}/{cat_name}: {len(docs)} documents"
                )

                for doc in docs:
                    if sample_mode and total_docs >= sample_size:
                        break
                    doc["year"] = year
                    doc["case_type"] = cat_name
                    yield doc
                    total_docs += 1

        logger.info(f"Total documents enumerated: {total_docs}")

    def fetch_all(self) -> Generator[dict, None, None]:
        """Yield all Constitutional Court decisions with full text."""
        seen_unids = set()

        for doc_info in self._enumerate_all_documents():
            unid = doc_info["unid"]
            if unid in seen_unids:
                continue
            seen_unids.add(unid)

            try:
                # Fetch decision page for metadata and file URLs
                meta = self._fetch_decision_page(unid)
                # Merge inline metadata
                for key in [
                    "signatura",
                    "date_str",
                    "decision_type",
                    "outcome",
                    "year",
                    "case_type",
                ]:
                    if key in doc_info and key not in meta:
                        meta[key] = doc_info[key]

                # Download full text
                full_text = self._download_full_text(meta)
                if not full_text:
                    logger.warning(
                        f"No full text for {meta.get('signatura', unid)}"
                    )
                    continue

                meta["full_text"] = full_text
                yield meta

            except Exception as e:
                logger.warning(
                    f"Failed to process {doc_info.get('signatura', unid)}: {e}"
                )

    def fetch_updates(self, since: datetime) -> Generator[dict, None, None]:
        """Yield documents from the most recent year(s) only."""
        seen_unids = set()
        current_year = datetime.now().year

        page_html = self._get_session()
        years = self._enumerate_years(page_html)

        # Only process years >= since.year
        for year, year_trigger in years:
            if year < since.year:
                break

            logger.info(f"Checking year {year} for updates...")
            try:
                year_html = self._xsp_post(year_trigger)
            except Exception as e:
                logger.warning(f"Failed to expand year {year}: {e}")
                continue

            idx_match = re.search(r'rpView:(\d+):', year_trigger)
            year_idx = idx_match.group(1) if idx_match else "0"

            subcats = self._enumerate_subcategories(year_html, year_idx)

            for cat_name, cat_trigger in subcats:
                try:
                    cat_html = self._xsp_post(cat_trigger)
                except Exception as e:
                    logger.warning(
                        f"Failed to expand {year}/{cat_name}: {e}"
                    )
                    continue

                docs = self._extract_document_links(cat_html)

                for doc_info in docs:
                    unid = doc_info["unid"]
                    if unid in seen_unids:
                        continue
                    seen_unids.add(unid)

                    # Check date
                    date_str = doc_info.get("date_str", "")
                    if date_str:
                        try:
                            parts = date_str.rstrip(".").split(".")
                            if len(parts) >= 3:
                                d = datetime(
                                    int(parts[2]),
                                    int(parts[1]),
                                    int(parts[0]),
                                    tzinfo=timezone.utc,
                                )
                                if d < since:
                                    continue
                        except (ValueError, IndexError):
                            pass

                    try:
                        meta = self._fetch_decision_page(unid)
                        for key in doc_info:
                            if key not in meta:
                                meta[key] = doc_info[key]

                        full_text = self._download_full_text(meta)
                        if full_text:
                            meta["full_text"] = full_text
                            yield meta
                    except Exception as e:
                        logger.warning(f"Failed to process {unid}: {e}")

    def normalize(self, raw: dict) -> dict:
        """Transform raw document data into standard schema."""
        unid = raw.get("unid", "")
        signatura = raw.get("signatura", "")
        full_text = raw.get("full_text", "")
        decision_type = raw.get("decision_type", "")
        outcome = raw.get("outcome", "")
        summary = raw.get("summary", "")

        # Parse date
        date_str = raw.get("decision_date", raw.get("date_str", ""))
        date_iso = ""
        if date_str:
            try:
                parts = date_str.rstrip(".").split(".")
                if len(parts) >= 3:
                    day = int(parts[0])
                    month = int(parts[1])
                    year = int(parts[2])
                    date_iso = f"{year:04d}-{month:02d}-{day:02d}"
            except (ValueError, IndexError):
                date_iso = date_str

        # Build title
        title = signatura
        if decision_type:
            title = f"{signatura} - {decision_type}"
        if outcome:
            title = f"{title} ({outcome})"

        url = f"{BASE_URL}{DECISION_PATH}?action=openDocument&documentId={unid}"

        return {
            "_id": f"HR-USUD-{signatura}" if signatura else f"HR-USUD-{unid}",
            "_source": "HR/ConstitutionalCourt",
            "_type": "case_law",
            "_fetched_at": datetime.now(timezone.utc).isoformat(),
            "title": title,
            "text": full_text,
            "date": date_iso,
            "url": url,
            "signatura": signatura,
            "decision_type": decision_type,
            "outcome": outcome,
            "summary": summary,
            "case_type": raw.get("case_type", ""),
            "year": raw.get("year", ""),
            "language": "hrv",
        }

    def test_connection(self):
        """Quick connectivity test."""
        print("Testing Croatian Constitutional Court (sljeme.usud.hr)...")

        print("\n1. Testing browse page...")
        try:
            page_html = self._get_session()
            print(f"   Status: OK, viewid obtained")

            years = self._enumerate_years(page_html)
            print(f"   Years found: {len(years)}")
            print(f"   Year range: {years[-1][0]}-{years[0][0]}")
        except Exception as e:
            print(f"   ERROR: {e}")
            return

        print("\n2. Expanding most recent year...")
        try:
            year, trigger = years[0]
            year_html = self._xsp_post(trigger)
            idx = re.search(r'rpView:(\d+):', trigger).group(1)
            subcats = self._enumerate_subcategories(year_html, idx)
            print(f"   Year {year}: {len(subcats)} case types")
            print(
                f"   Types: {', '.join(s[0] for s in subcats[:10])}"
            )
        except Exception as e:
            print(f"   ERROR: {e}")
            return

        print("\n3. Expanding first subcategory...")
        try:
            cat_name, cat_trigger = subcats[0]
            cat_html = self._xsp_post(cat_trigger)
            docs = self._extract_document_links(cat_html)
            print(f"   {cat_name}: {len(docs)} documents")
            if docs:
                print(f"   Sample: {docs[0].get('link_text', 'N/A')}")
        except Exception as e:
            print(f"   ERROR: {e}")
            return

        print("\n4. Fetching a decision...")
        if docs:
            try:
                unid = docs[0]["unid"]
                meta = self._fetch_decision_page(unid)
                print(f"   Signatura: {meta.get('signatura', 'N/A')}")
                print(f"   Date: {meta.get('decision_date', 'N/A')}")
                print(f"   DOCX URLs: {len(meta.get('docx_urls', []))}")
                print(f"   PDF URLs: {len(meta.get('pdf_urls', []))}")

                full_text = self._download_full_text(meta)
                print(f"   Full text: {len(full_text)} chars")
                if full_text:
                    print(f"   Preview: {full_text[:150]}...")
            except Exception as e:
                print(f"   ERROR: {e}")

        print("\nTest complete!")


def main():
    scraper = CroatianConstitutionalCourtScraper()

    if len(sys.argv) < 2:
        print(
            "Usage: python bootstrap.py [bootstrap|update|test] "
            "[--sample] [--sample-size N]"
        )
        sys.exit(1)

    command = sys.argv[1]
    sample_mode = "--sample" in sys.argv
    sample_size = 12
    if "--sample-size" in sys.argv:
        idx = sys.argv.index("--sample-size")
        sample_size = int(sys.argv[idx + 1])

    if command == "test":
        scraper.test_connection()

    elif command == "bootstrap":
        if sample_mode:
            stats = scraper.run_sample(n=sample_size)
            print(
                f"\nSample complete: "
                f"{stats.get('sample_records_saved', 0)} records saved to sample/"
            )
        else:
            stats = scraper.bootstrap()
            print(
                f"\nBootstrap complete: {stats['records_new']} new, "
                f"{stats['records_updated']} updated, "
                f"{stats['records_skipped']} skipped"
            )
        print(json.dumps(stats, indent=2))

    elif command == "update":
        stats = scraper.update()
        print(
            f"\nUpdate complete: {stats['records_new']} new, "
            f"{stats['records_updated']} updated"
        )
        print(json.dumps(stats, indent=2))

    else:
        print(f"Unknown command: {command}")
        sys.exit(1)


if __name__ == "__main__":
    main()
